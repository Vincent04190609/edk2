#!/usr/bin/env python3
"""Update Engineering Spec.docx for BIOS feature changes."""

from __future__ import annotations

import argparse
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document

from spec_paths import engineering_spec_path


def find_paragraph(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    return None


def bump_metadata(doc: Document) -> tuple[str, str]:
    version_paragraph = find_paragraph(doc, "Version:")
    date_paragraph = find_paragraph(doc, "Date:")
    if version_paragraph is None or date_paragraph is None:
        raise ValueError("Version: or Date: paragraph not found in Engineering Spec")

    version_text = version_paragraph.text.split(":", 1)[1].strip()
    match = re.match(r"^(\d+)\.(\d+)$", version_text)
    if match:
        major, minor = match.groups()
        new_version = f"{major}.{int(minor) + 1}"
    else:
        new_version = version_text

    new_date = datetime.now().strftime("%b. %d. %Y")
    version_paragraph.text = f"Version: {new_version}"
    date_paragraph.text = f"Date: {new_date}"
    return new_version, new_date


def find_table_by_headers(doc: Document, headers: tuple[str, ...]) -> int | None:
    for index, table in enumerate(doc.tables):
        if not table.rows:
            continue
        row_values = [cell.text.strip() for cell in table.rows[0].cells]
        if row_values[: len(headers)] == list(headers):
            return index
    return None


def upsert_table_row(table, key_col: int, value_col: int, name: str, value: str) -> str:
    name_lower = name.strip().lower()
    for row in table.rows[1:]:
        if row.cells[key_col].text.strip().lower() == name_lower:
            row.cells[value_col].text = value
            return "updated"
    row = table.add_row()
    row.cells[key_col].text = name
    row.cells[value_col].text = value
    return "added"


def update_smbios_field(doc: Document, field_name: str, field_value: str) -> str:
    table_index = find_table_by_headers(doc, ("Name", "Default Value"))
    if table_index is None:
        raise ValueError("SMBIOS table with headers Name / Default Value not found")
    action = upsert_table_row(doc.tables[table_index], 0, 1, field_name, field_value)
    return f"smbios_field:{action}:{field_name}={field_value}"


def update_setup_item(doc: Document, item_name: str, current_setting: str) -> str:
    table_index = find_table_by_headers(doc, ("Items", "Current Setting"))
    if table_index is None:
        raise ValueError("Setup table with headers Items / Current Setting not found")
    action = upsert_table_row(doc.tables[table_index], 0, 1, item_name, current_setting)
    return f"setup_item:{action}:{item_name}={current_setting}"


def main() -> int:
    parser = argparse.ArgumentParser(description="Update Engineering Spec.docx")
    parser.add_argument("--path", help="Override Engineering Spec path")
    parser.add_argument(
        "--bump-metadata",
        action="store_true",
        help="Increment Version x.y and set Date to today",
    )
    parser.add_argument("--smbios-field", nargs=2, metavar=("NAME", "VALUE"))
    parser.add_argument("--setup-item", nargs=2, metavar=("ITEM", "SETTING"))
    args = parser.parse_args()

    if not any([args.bump_metadata, args.smbios_field, args.setup_item]):
        parser.error("Specify at least one update action")

    try:
        spec_path = engineering_spec_path(args.path)
    except Exception as exc:
        print(str(exc), file=sys.stderr)
        return 1

    if not spec_path.is_file():
        print(
            f"Engineering Spec not found: {spec_path}\n"
            "Hint: set mode to 'windows' in .cursor/kb-path-config.json on native Windows, "
            "or ensure Fleet bind mounts are active for fleet mode.",
            file=sys.stderr,
        )
        return 1

    doc = Document(str(spec_path))
    changes: list[str] = []

    if args.bump_metadata:
        version, date = bump_metadata(doc)
        changes.append(f"metadata:version={version},date={date}")

    if args.smbios_field:
        name, value = args.smbios_field
        changes.append(update_smbios_field(doc, name, value))

    if args.setup_item:
        item, setting = args.setup_item
        changes.append(update_setup_item(doc, item, setting))

    doc.save(str(spec_path))
    print("\n".join(changes))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
